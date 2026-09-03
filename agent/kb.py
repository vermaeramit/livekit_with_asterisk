"""
Knowledge base: layout-aware PDF ingestion + hybrid retrieval.

Extraction uses pymupdf4llm, not plain get_text(). On designed layouts
(quotes, brochures, multi-column policy docs) raw text extraction reads across
columns instead of down them, which silently fuses unrelated content - e.g. two
pricing options merging into one unusable chunk. pymupdf4llm resolves reading
order and emits markdown, giving real headings and tables to chunk on.

ingest_file() is the single entry point: CLI today, folder watcher and the
Step 11 admin UI later.
"""
from __future__ import annotations

import asyncio
import hashlib
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path

import pymupdf4llm
import tiktoken
from openai import AsyncOpenAI

import store
import webkb

log = logging.getLogger("kb")

EMBED_MODEL = os.getenv("EMBED_MODEL", "text-embedding-3-small")
TARGET_TOKENS = int(os.getenv("KB_CHUNK_TOKENS", "250"))
OVERLAP_TOKENS = int(os.getenv("KB_CHUNK_OVERLAP", "50"))
MIN_CHUNK_TOKENS = int(os.getenv("KB_MIN_CHUNK_TOKENS", "40"))
LEX_THRESHOLD = float(os.getenv("KB_LEX_THRESHOLD", "0.20"))
EMBED_BATCH = 96

try:
    _enc = tiktoken.get_encoding("cl100k_base")
except Exception:  # pragma: no cover
    _enc = None


def ntok(s: str) -> int:
    return len(_enc.encode(s)) if _enc else max(1, len(s) // 4)


# ────────────────────────────── extraction ──────────────────────────────

_PIC = re.compile(r"<!--\s*(?:Start|End) of picture text\s*-->")
_TAG = re.compile(r"</?(?:sup|sub|span|div)[^>]*>")


def clean_md(md: str) -> str:
    """Strip pymupdf4llm's markup scaffolding, keep the text it wraps.

    Picture-text markers wrap real content (logos with taglines, badge graphics),
    so the markers go but the text stays.
    """
    md = _PIC.sub("", md)
    md = md.replace("<br>", "\n").replace("<br/>", "\n")
    md = _TAG.sub("", md)
    md = re.sub(r"[ \t]+\n", "\n", md)
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def extract_pdf(path: str) -> tuple[list[tuple[int, str]], int]:
    """-> ([(page_no, markdown)], page_count)"""
    pages = pymupdf4llm.to_markdown(path, page_chunks=True, show_progress=False)
    out = []
    for p in pages:
        md = clean_md(p.get("text", ""))
        if md:
            out.append((p.get("metadata", {}).get("page", len(out) + 1), md))
    return out, len(pages)


def _docx_table_md(table) -> str:
    """A docx table as a markdown pipe table.

    Emitted as markdown rather than flattened to prose because the chunker
    recognises pipe tables and keeps them whole - a price table split across
    two chunks retrieves as two half-answers.

    The first row is treated as the header. That is wrong for a table with no
    header, but a markdown table must have one, and losing a data row to the
    header beats emitting something no parser reads as a table at all.
    """
    rows = []
    for row in table.rows:
        # A merged cell appears once per grid position in python-docx, so a
        # horizontally merged one repeats its text. Left alone: collapsing it
        # would mean guessing which column the value really belongs to.
        cells = [c.text.replace("\n", " ").replace("|", r"\|").strip()
                 for c in row.cells]
        if any(cells):
            rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    header, *body = rows
    sep = "| " + " | ".join("---" for _ in table.columns) + " |"
    return "\n".join([header, sep, *body])


def extract_docx(path: str) -> tuple[list[tuple[int, str]], int]:
    """-> ([(None, markdown)], 0) - a Word document as one markdown stream.

    Page is None, not 1. A .docx has no pages until something lays it out, and
    the number would differ between Word, LibreOffice and a different paper
    size. The console hides a null page rather than printing "p.1" on every
    chunk of every document, and the heading path is the real locator anyway.

    Walks the document BODY in order rather than iterating .paragraphs and
    .tables separately. Those are two independent lists, so a table sitting
    between two headings comes out after every paragraph in the document -
    silently detaching every table from the text that explains it.
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(path)
    parts: list[str] = []

    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]

        if tag == "tbl":
            md = _docx_table_md(Table(child, doc))
            if md:
                parts.append(md)
            continue
        if tag != "p":
            continue

        para = Paragraph(child, doc)
        text = para.text.strip()
        if not text:
            continue

        style = (para.style.name or "") if para.style is not None else ""

        # "Heading 2" -> "## ". This is what the chunker splits on, and it is
        # why Word suits this pipeline: the structure is already in the file.
        m = _DOCX_HEADING.match(style)
        if m:
            parts.append("#" * min(int(m.group(1)), 6) + " " + text)
        elif style == "Title":
            parts.append("# " + text)
        elif style == "Subtitle":
            parts.append("## " + text)
        elif "List" in style:
            parts.append("- " + text)
        else:
            parts.append(text)

    md = clean_md("\n\n".join(parts))
    return ([(None, md)] if md else []), 0


# Extension -> extractor. Everything downstream works on [(page, markdown)],
# which is why adding a format is one function rather than a second pipeline.
EXTRACTORS = {".pdf": extract_pdf, ".docx": extract_docx}
SUPPORTED = tuple(EXTRACTORS)


# ────────────────────────────── chunking ──────────────────────────────

@dataclass
class Chunk:
    seq: int
    page: int
    heading: str
    content: str
    n_tokens: int = 0
    embed_text: str = field(default="", repr=False)


_H = re.compile(r"^(#{1,6})\s+(.+?)\s*#*$")
# Word's own heading styles, mapped onto the markdown ones above.
_DOCX_HEADING = re.compile(r"^Heading (\d+)$")
_TABLE_ROW = re.compile(r"^\s*\|.*\|\s*$")


def _strip_md(s: str) -> str:
    return re.sub(r"[*_`]", "", s).strip()


def chunk_markdown(pages: list[tuple[int, str]], title: str) -> list[Chunk]:
    """Split on markdown headings, then pack to TARGET_TOKENS with overlap.

    Heading depth is tracked as a stack, so a chunk carries its full path
    ("DAY-BY-DAY ITINERARY > Desert Safari with BBQ Dinner"). Short questions
    retrieve noticeably better when a chunk knows where it came from.
    """
    chunks: list[Chunk] = []
    stack: list[tuple[int, str]] = []
    buf: list[str] = []
    buf_tok = 0
    cur_page = pages[0][0] if pages else 1

    def heading_path() -> str:
        return " > ".join(h for _, h in stack)

    def emit():
        nonlocal buf, buf_tok
        content = "\n".join(buf).strip()
        buf, buf_tok = [], 0
        if len(content) < 25:
            return
        path = " > ".join(x for x in (title, heading_path()) if x)
        chunks.append(Chunk(
            seq=len(chunks), page=cur_page, heading=heading_path(),
            content=content, n_tokens=ntok(content),
            embed_text=f"{path}\n\n{content}" if path else content,
        ))

    for page_no, md in pages:
        cur_page = page_no
        lines = md.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            m = _H.match(line)
            if m:
                emit()
                depth, text = len(m.group(1)), _strip_md(m.group(2))
                while stack and stack[-1][0] >= depth:
                    stack.pop()
                if text:
                    stack.append((depth, text))
                i += 1
                continue

            # a markdown table is atomic - splitting one destroys the column
            # alignment and the rows become meaningless
            if _TABLE_ROW.match(line):
                tbl = []
                while i < len(lines) and (_TABLE_ROW.match(lines[i]) or not lines[i].strip()):
                    if lines[i].strip():
                        tbl.append(lines[i])
                    i += 1
                block = "\n".join(tbl)
                if buf and buf_tok + ntok(block) > TARGET_TOKENS:
                    emit()
                buf.append(block)
                buf_tok += ntok(block)
                continue

            if not line.strip():
                i += 1
                continue

            t = ntok(line)
            if buf and buf_tok + t > TARGET_TOKENS:
                tail, tail_tok = [], 0
                for prev in reversed(buf):
                    pt = ntok(prev)
                    if tail_tok + pt > OVERLAP_TOKENS:
                        break
                    tail.insert(0, prev)
                    tail_tok += pt
                emit()
                buf, buf_tok = tail, tail_tok
            buf.append(line)
            buf_tok += t
            i += 1
    emit()
    return _absorb_tiny(chunks)


def _absorb_tiny(chunks: list[Chunk]) -> list[Chunk]:
    """Fold sub-threshold chunks into a neighbour.

    A 5-token chunk answers nothing but still competes in retrieval and
    displaces a real result. Merge forward (keeps it under its own heading),
    else backward.
    """
    out: list[Chunk] = []
    i = 0
    while i < len(chunks):
        c = chunks[i]
        if c.n_tokens < MIN_CHUNK_TOKENS:
            if i + 1 < len(chunks):
                nxt = chunks[i + 1]
                nxt.content = c.content + "\n" + nxt.content
                nxt.embed_text = c.content + "\n" + nxt.embed_text
                nxt.n_tokens = ntok(nxt.content)
                i += 1
                continue
            if out:
                out[-1].content += "\n" + c.content
                out[-1].embed_text += "\n" + c.content
                out[-1].n_tokens = ntok(out[-1].content)
                i += 1
                continue
        out.append(c)
        i += 1
    for n, c in enumerate(out):
        c.seq = n
    return out


# ────────────────────────────── embeddings ──────────────────────────────

# One client per key, not one client overall. Embedding is billed to whoever
# owns the documents, so an ingest for client A and a search for client B run on
# different credentials inside the same process - a single global client would
# quietly put both on whichever key happened to be first.
_clients: dict[str, AsyncOpenAI] = {}


def _openai(api_key: str | None = None) -> AsyncOpenAI:
    # Falling back to the platform key keeps the CLI and the cache warmer
    # working; neither belongs to a client.
    key = api_key or os.environ["OPENAI_API_KEY"]
    if key not in _clients:
        _clients[key] = AsyncOpenAI(api_key=key)
    return _clients[key]


async def embed(texts: list[str], on_batch=None,
                api_key: str | None = None) -> list[list[float]]:
    """on_batch(done, total) is awaited after each batch, if given.

    Embedding is the long pole of an ingest and the only stage with a real
    denominator, which makes it the one worth reporting.
    """
    out: list[list[float]] = []
    for i in range(0, len(texts), EMBED_BATCH):
        r = await _openai(api_key).embeddings.create(
            model=EMBED_MODEL, input=texts[i:i + EMBED_BATCH])
        out.extend(d.embedding for d in sorted(r.data, key=lambda d: d.index))
        if on_batch:
            await on_batch(len(out), len(texts))
    return out


def _vec(v: list[float]) -> str:
    """asyncpg has no pgvector codec - send a literal and cast in SQL."""
    return "[" + ",".join(f"{x:.7g}" for x in v) + "]"


# ────────────────────────────── ingest ──────────────────────────────

async def ingest_file(path: str, config_name: str = "default",
                      force: bool = False, campaign_id: int | None = None,
                      on_progress=None, api_key: str | None = None) -> dict:
    """on_progress(event: dict) is awaited at each stage, if given.

    Stages: hashing, extracting, chunking, embedding (with done/total), saving.
    The CLI passes nothing and behaves exactly as before.
    """
    async def emit(**event):
        if on_progress:
            await on_progress(event)

    p = Path(path)
    await emit(stage="hashing")
    digest = hashlib.sha256(p.read_bytes()).hexdigest()
    pool = await store.pool()

    existing = await pool.fetchrow(
        "SELECT id, content_hash, chunk_count FROM kb_documents "
        "WHERE config_name=$1 AND filename=$2", config_name, p.name)
    if existing and existing["content_hash"] == digest and not force:
        return {"file": p.name, "status": "unchanged", "chunks": existing["chunk_count"]}

    # pymupdf4llm and the chunker are CPU-bound and hold the GIL for seconds on a
    # 50-page document. On the CLI that only makes the prompt wait; inside the
    # admin API it would freeze every other request for the length of an upload.
    await emit(stage="extracting")
    extractor = EXTRACTORS.get(p.suffix.lower())
    if extractor is None:
        return {"file": p.name, "status": "empty",
                "error": f"{p.suffix or 'this file'} is not supported - "
                         f"use {' or '.join(SUPPORTED)}"}
    pages, n_pages = await asyncio.to_thread(extractor, str(p))
    if not pages:
        # The two formats fail emptily for different reasons, and the difference
        # decides what the person holding the file should do next.
        why = ("no extractable text - scanned PDF? OCR would be needed"
               if p.suffix.lower() == ".pdf"
               else "the document has no text - images and text boxes are not read")
        return {"file": p.name, "status": "empty", "error": why}

    return await _chunk_embed_store(
        pages, n_pages, p.name, digest, config_name, campaign_id, api_key,
        emit, existing, source_id=None, source_url=None)


async def _chunk_embed_store(pages, n_pages, filename, digest, config_name,
                             campaign_id, api_key, emit, existing,
                             source_id=None, source_url=None):
    """Chunk, embed and save one document, whether it came from a file or a URL.

    Shared rather than copied. The two callers differ only in how they got hold
    of the text, and a second copy of this is a second place to remember when
    the chunker or the schema changes - which is the failure this codebase has
    already had four times with JSONB columns.
    """
    pool = await store.pool()
    title = _title_of(filename)
    await emit(stage="chunking", pages=n_pages)
    chunks = await asyncio.to_thread(chunk_markdown, pages, title)
    if not chunks:
        return {"file": filename, "status": "empty", "error": "no chunks produced"}

    await emit(stage="embedding", done=0, total=len(chunks))
    vectors = await embed(
        [c.embed_text for c in chunks],
        on_batch=lambda done, total: emit(stage="embedding", done=done, total=total),
        api_key=api_key,
    )
    await emit(stage="saving", chunks=len(chunks))

    # one transaction: a failed ingest leaves the previous version intact
    # rather than a half-updated KB
    async with pool.acquire() as conn:
        async with conn.transaction():
            if existing:
                doc_id = existing["id"]
                await conn.execute("DELETE FROM kb_chunks WHERE doc_id=$1", doc_id)
                # COALESCE so a CLI re-ingest (which passes no campaign_id) does
                # not orphan a document the panel had already scoped to a tenant
                await conn.execute(
                    "UPDATE kb_documents SET content_hash=$2, page_count=$3, "
                    "chunk_count=$4, title=$5, campaign_id=COALESCE($6, campaign_id), "
                    "source_id=COALESCE($7, source_id), "
                    "source_url=COALESCE($8, source_url), "
                    "updated_at=now() WHERE id=$1",
                    doc_id, digest, n_pages, len(chunks), title, campaign_id,
                    source_id, source_url)
            else:
                doc_id = await conn.fetchval(
                    "INSERT INTO kb_documents (config_name, filename, title, "
                    "content_hash, page_count, chunk_count, campaign_id, "
                    "source_id, source_url) "
                    "VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9) RETURNING id",
                    config_name, filename, title, digest, n_pages, len(chunks),
                    campaign_id, source_id, source_url)
            await conn.executemany(
                "INSERT INTO kb_chunks (doc_id, config_name, campaign_id, seq, page, "
                "heading, content, n_tokens, embedding) "
                "VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9::vector)",
                [(doc_id, config_name, campaign_id, c.seq, c.page, c.heading,
                  c.content, c.n_tokens, _vec(v)) for c, v in zip(chunks, vectors)])

    return {"file": filename, "status": "updated" if existing else "created",
            "pages": n_pages, "chunks": len(chunks),
            "tokens": sum(c.n_tokens for c in chunks)}


def _title_of(filename: str) -> str:
    stem = filename.rsplit(".", 1)[0]
    return stem.replace("_", " ").replace("-", " ")


async def ingest_web_source(url: str, config_name: str = "default",
                            campaign_id: int | None = None,
                            source_id: int | None = None,
                            api_key: str | None = None,
                            emit=None, force: bool = False) -> dict:
    """Import a URL, or every sheet of a workbook published at one.

    One kb_document per page, so refresh, disable and citations all work per
    sheet exactly as they do per file. A workbook of 47 tabs becomes 47
    documents, and the one holding 2000 vendor rows can be switched off
    without touching the other 46.
    """
    emit = emit or (lambda **_: asyncio.sleep(0))
    pool = await store.pool()

    await emit(stage="fetching", url=url)
    final_url, ctype, body = await asyncio.to_thread(webkb.fetch, url)

    # A URL that serves a PDF is a PDF. Nothing about it needs a web importer.
    if "pdf" in ctype.lower() or final_url.lower().endswith(".pdf"):
        pages_out = [(final_url, url.rsplit("/", 1)[-1] or "document.pdf")]
    else:
        markup = webkb._decode(body, ctype)
        if webkb.is_workbook(markup):
            pages_out = webkb.workbook_pages(final_url, markup)
            if not pages_out:
                return {"url": url, "status": "empty",
                        "error": "this looks like a workbook but its sheet list "
                                 "could not be read"}
        else:
            pages_out = [(final_url, webkb.page_title(markup, "page"))]

    results, skipped = [], []
    for i, (page_url, name) in enumerate(pages_out):
        await emit(stage="page", done=i, total=len(pages_out), name=name)
        try:
            _, page_ctype, page_body = await asyncio.to_thread(webkb.fetch, page_url)
        except Exception as e:
            skipped.append({"name": name, "why": f"{type(e).__name__}: {e}"})
            continue

        text, images = webkb.extract(webkb._decode(page_body, page_ctype))
        if webkb.prose_length(text) < webkb.MIN_TEXT_CHARS:
            # A sheet of screenshots. Reported rather than dropped silently:
            # otherwise nobody learns that the oil price list is a picture and
            # the agent simply does not know it.
            skipped.append({"name": name, "why": f"no readable text ({images} images)"})
            continue

        # The name is the identity. A sheet renamed upstream becomes a new
        # document and the old one is removed on the next refresh, which is
        # right - a renamed tab is usually a different tab.
        filename = _web_filename(name)
        digest = hashlib.sha256(text.encode()).hexdigest()
        existing = await pool.fetchrow(
            "SELECT id, content_hash, chunk_count FROM kb_documents "
            "WHERE config_name=$1 AND filename=$2", config_name, filename)
        if existing and existing["content_hash"] == digest and not force:
            results.append({"file": filename, "status": "unchanged",
                            "chunks": existing["chunk_count"]})
            continue

        results.append(await _chunk_embed_store(
            [(None, text)], 1, filename, digest, config_name, campaign_id,
            api_key, emit, existing, source_id=source_id, source_url=page_url))

    # Anything this source produced last time and does not produce now is gone
    # upstream. Left behind, the agent would keep answering from a sheet that
    # no longer exists - an offer that ended, a price that was withdrawn.
    removed = 0
    if source_id is not None:
        keep = [r["file"] for r in results]
        removed = await pool.fetchval(
            "WITH gone AS (DELETE FROM kb_documents WHERE source_id=$1 "
            "  AND NOT (filename = ANY($2::text[])) RETURNING 1) "
            "SELECT count(*) FROM gone", source_id, keep) or 0

    return {"url": url, "status": "ok", "pages": len(results),
            "removed": removed, "documents": results, "skipped": skipped}


def _web_filename(name: str) -> str:
    """A stable identity for a page, safe as a filename and as a citation."""
    safe = re.sub(r"[^\w .&+-]", " ", name).strip()
    safe = re.sub(r"\s+", " ", safe)[:120] or "page"
    return f"{safe}.web"


# ────────────────────────────── retrieval ──────────────────────────────

async def search(query: str, config_name: str = "default",
                 top_k: int = 3, min_score: float = 0.25,
                 api_key: str | None = None) -> list[dict]:
    """Hybrid: cosine similarity + trigram word similarity.

    word_similarity, NOT similarity: the latter compares whole strings, so a
    250-token chunk against a 6-word question scores near zero and the lexical
    leg never fires. word_similarity asks how well the query matches SOME PART
    of the chunk - which is the actual question.
    """
    if not query.strip():
        return []
    qvec = _vec((await embed([query], api_key=api_key))[0])
    rows = await (await store.pool()).fetch(
        """
        WITH vec AS (
            SELECT id, doc_id, page, heading, content,
                   1 - (embedding <=> $1::vector) AS score, 'vec' AS src
              FROM kb_chunks
             WHERE config_name = $2 AND embedding IS NOT NULL
             ORDER BY embedding <=> $1::vector
             LIMIT $3
        ), lex AS (
            SELECT id, doc_id, page, heading, content,
                   word_similarity($4, content) AS score, 'lex' AS src
              FROM kb_chunks
             WHERE config_name = $2 AND word_similarity($4, content) > $5
             ORDER BY word_similarity($4, content) DESC
             LIMIT $3
        ), merged AS (SELECT * FROM vec UNION ALL SELECT * FROM lex)
        SELECT DISTINCT ON (id) id, doc_id, page, heading, content, score, src
          FROM merged ORDER BY id, score DESC
        """,
        qvec, config_name, top_k * 2, query, LEX_THRESHOLD)

    hits = [dict(r) for r in rows if r["score"] is not None and r["score"] >= min_score]
    hits.sort(key=lambda h: h["score"], reverse=True)
    return hits[:top_k]


def format_context(hits: list[dict]) -> str:
    if not hits:
        return ""
    return "\n\n".join(
        f"[{i}]{' (' + h['heading'] + ')' if h.get('heading') else ''}\n{h['content']}"
        for i, h in enumerate(hits, 1))


# ────────────────────────────── inline context ──────────────────────────────

async def load_inline(config_name: str = "default",
                      max_tokens: int = 6000) -> tuple[str, int, str]:
    """Build the always-in-prompt knowledge layer.

    Small KBs go into the prompt whole. That removes the per-turn embedding call
    (measured 390-1244 ms) AND the whole class of retrieval misses - the model
    sees everything and picks for itself.

    Above the budget we fall back to an index of document titles and headings.
    The model still needs to know WHAT is searchable, otherwise it either never
    calls the tool or calls it on every turn.

    -> (text, n_tokens, mode)   mode is "full" | "index" | "empty"
    """
    rows = await (await store.pool()).fetch(
        """SELECT d.title, c.heading, c.content, c.n_tokens
             FROM kb_chunks c JOIN kb_documents d ON d.id = c.doc_id
            WHERE c.config_name = $1 AND d.enabled
            ORDER BY d.filename, c.seq""", config_name)
    if not rows:
        return "", 0, "empty"

    total = sum(r["n_tokens"] or 0 for r in rows)
    if total <= max_tokens:
        parts, last = [], None
        for r in rows:
            if r["title"] != last:
                parts.append(f"\n## {r['title']}")
                last = r["title"]
            head = f"### {r['heading']}\n" if r["heading"] else ""
            parts.append(f"{head}{r['content']}")
        return "\n\n".join(parts).strip(), total, "full"

    seen, lines, last = set(), [], None
    for r in rows:
        if r["title"] != last:
            lines.append(f"\n## {r['title']}")
            last = r["title"]
        h = r["heading"]
        if h and h not in seen:
            seen.add(h)
            lines.append(f"  - {h}")
    idx = "\n".join(lines).strip()
    return idx, ntok(idx), "index"
